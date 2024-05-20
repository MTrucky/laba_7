import kivy
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.properties import StringProperty, NumericProperty
from kivy.uix.popup import Popup
from abc import ABC, abstractmethod
import docx
from openpyxl import Workbook

class Building(ABC):
    def __init__(self, length, width, height):
        self.length = length
        self.width = width
        self.height = height

    @abstractmethod
    def calculate_area(self):
        pass

    @abstractmethod
    def calculate_heat_power(self):
        pass

    def __str__(self):
        return f"Building: {self.length} x {self.width} x {self.height}"

    def __repr__(self):
        return f"<Building(length={self.length}, width={self.width}, height={self.height})>"

class Room(Building):
    def calculate_area(self):
        return self.length * self.width

    def calculate_heat_power(self):
        return self.calculate_area() * self.height * 100  # Примерно

    def __str__(self):
        return f"Room: {self.length} x {self.width} x {self.height}"

    def __repr__(self):
        return f"<Room(length={self.length}, width={self.width}, height={self.height})>"

class Apartment(Building):
    def __init__(self, length, width, height, num_rooms):
        super().__init__(length, width, height)
        self.num_rooms = num_rooms

    def calculate_area(self):
        return self.length * self.width

    def calculate_total_area(self):
        return self.calculate_area() * self.num_rooms

    def calculate_heat_power(self):
        return self.calculate_total_area() * 80  # Примерно

    def __str__(self):
        return f"Apartment: {self.length} x {self.width} x {self.height}, Rooms: {self.num_rooms}"

    def __repr__(self):
        return f"<Apartment(length={self.length}, width={self.width}, height={self.height}, num_rooms={self.num_rooms})>"

class MultistoryBuilding(Building):
    def __init__(self, length, width, height, num_floors, num_units_per_floor):
        super().__init__(length, width, height)
        self.num_floors = num_floors
        self.num_units_per_floor = num_units_per_floor

    def calculate_area(self):
        return self.length * self.width

    def calculate_total_area(self):
        return self.calculate_area() * self.num_units_per_floor * self.num_floors

    def calculate_heat_power(self):
        return self.calculate_total_area() * 70  # Примерно

    def __str__(self):
        return f"MultistoryBuilding: {self.length} x {self.width} x {self.height}, Floors: {self.num_floors}, Units per floor: {self.num_units_per_floor}"

    def __repr__(self):
        return f"<MultistoryBuilding(length={self.length}, width={self.width}, height={self.height}, num_floors={self.num_floors}, num_units_per_floor={self.num_units_per_floor})>"

class BuildingApp(App):
    building_type = StringProperty("Комната")
    length = NumericProperty(0.0)
    width = NumericProperty(0.0)
    height = NumericProperty(0.0)
    num_rooms = NumericProperty(0)
    num_floors = NumericProperty(0)
    num_units_per_floor = NumericProperty(0)

    def build(self):
        self.title = "Калькулятор строительства"
        layout = BoxLayout(orientation='vertical')

        building_type_label = Label(text="Выберите тип помещения:")
        layout.add_widget(building_type_label)

        self.building_type_spinner = Spinner(
            text=self.building_type,
            values=["Комната", "Квартира", "Многоэтажный дом"]
        )
        layout.add_widget(self.building_type_spinner)

        layout.add_widget(Label(text="Длина:"))
        self.length_input = TextInput(multiline=False)
        layout.add_widget(self.length_input)

        layout.add_widget(Label(text="Ширина:"))
        self.width_input = TextInput(multiline=False)
        layout.add_widget(self.width_input)

        layout.add_widget(Label(text="Высота:"))
        self.height_input = TextInput(multiline=False)
        layout.add_widget(self.height_input)

        layout.add_widget(Label(text="Количество комнат (если квартира):"))
        self.num_rooms_input = TextInput(multiline=False)
        layout.add_widget(self.num_rooms_input)

        layout.add_widget(Label(text="Этажей (если многоэтажный дом):"))
        self.num_floors_input = TextInput(multiline=False)
        layout.add_widget(self.num_floors_input)

        layout.add_widget(Label(text="Количество квартир на этаже (если многоэтажный дом):"))
        self.num_units_input = TextInput(multiline=False)
        layout.add_widget(self.num_units_input)

        calculate_button = Button(text="Рассчитать и сохранить", on_press=self.calculate_and_save_report)
        layout.add_widget(calculate_button)

        self.result_label = Label(text="")
        layout.add_widget(self.result_label)

        return layout

    def calculate_and_save_report(self, instance):
        building_type = self.building_type_spinner.text

        length = float(self.length_input.text)
        width = float(self.width_input.text)
        height = float(self.height_input.text)

        if building_type == "Комната":
            building = Room(length, width, height)
            total_area = building.calculate_area()
            heat_power = building.calculate_heat_power()
        elif building_type == "Квартира":
            num_rooms = int(self.num_rooms_input.text)
            building = Apartment(length, width, height, num_rooms)
            total_area = building.calculate_total_area()
            heat_power = building.calculate_heat_power()
        else:  # Многоэтажный дом
            num_floors = int(self.num_floors_input.text)
            num_units_per_floor = int(self.num_units_input.text)
            building = MultistoryBuilding(length, width, height, num_floors, num_units_per_floor)
            total_area = building.calculate_total_area()
            heat_power = building.calculate_heat_power()

        self.result_label.text = f"Общая площадь: {total_area} кв.м\nТепловая мощность: {heat_power} Вт"

        # Сохранение результатов в отчет .docx
        doc = docx.Document()
        doc.add_heading('Результаты расчетов', level=1)
        doc.add_paragraph(f"Общая площадь: {total_area} кв.м")
        doc.add_paragraph(f"Тепловая мощность: {heat_power} Вт")
        doc.save('report.docx')

        # Сохранение результатов в отчет .xlsx
        wb = Workbook()
        ws = wb.active
        ws.append(["Общая площадь", "Тепловая мощность"])
        ws.append([total_area, heat_power])
        wb.save('report.xlsx')

        popup = Popup(title='Сохранение',
                      content=Label(text='Результаты сохранены в report.docx и report.xlsx'),
                      size_hint=(None, None), size=(400, 200))
        popup.open()

if __name__ == '__main__':
    BuildingApp().run()
